import os
import fitz  # PyMuPDF
import docx  # python-docx
from services.text_processor import clean_text
from services.section_detector import detect_sections, CANONICAL_SECTIONS
from services.skill_extractor import extract_skills
from services.resume_analyzer import analyze_resume, CATEGORY_WEIGHTS
from services.job_matcher import match_resume_to_job
from services.recommendation_service import generate_recommendations
from services.evidence_analyzer import analyze_resume_evidence
from services.candidate_profile import build_candidate_profile

# Resource limit safeguard for PDF page count
MAX_PDF_PAGES = 50


def parse_pdf(file_path: str) -> dict:
    """
    Extracts text from a PDF file using PyMuPDF (fitz).
    
    Returns:
        dict: containing page_count, raw_text, and error if any.
    """
    page_texts = []
    page_count = 0

    try:
        with fitz.open(file_path) as doc:
            page_count = doc.page_count
            if doc.is_encrypted:
                return {
                    "page_count": page_count,
                    "raw_text": "",
                    "error": "Document is password protected or encrypted."
                }

            if page_count > MAX_PDF_PAGES:
                return {
                    "page_count": page_count,
                    "raw_text": "",
                    "error": f"PDF exceeds maximum allowed page limit ({MAX_PDF_PAGES} pages)."
                }

            for page in doc:
                text = page.get_text("text")
                if text:
                    page_texts.append(text)

        raw_text = "\n\n".join(page_texts)
        return {
            "page_count": page_count,
            "raw_text": raw_text,
            "error": None
        }
    except Exception as e:
        return {
            "page_count": page_count,
            "raw_text": "",
            "error": f"Failed to parse PDF file: {str(e)}"
        }


def parse_docx(file_path: str) -> dict:
    """
    Extracts text from a DOCX file using python-docx.
    Extracts both paragraphs and table text in document order.
    
    Returns:
        dict: containing page_count (None for DOCX), raw_text, and error if any.
    """
    extracted_chunks = []

    try:
        doc = docx.Document(file_path)
        
        for paragraph in doc.paragraphs:
            txt = paragraph.text.strip()
            if txt:
                extracted_chunks.append(txt)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    extracted_chunks.append(" | ".join(row_cells))

        raw_text = "\n\n".join(extracted_chunks)
        return {
            "page_count": None,
            "raw_text": raw_text,
            "error": None
        }
    except Exception as e:
        return {
            "page_count": None,
            "raw_text": "",
            "error": f"Failed to parse DOCX file: {str(e)}"
        }


def parse_resume(file_path: str, job_description: str = None) -> dict:
    """
    Main entry point for parsing and analyzing a resume file (PDF or DOCX).
    Optionally matches the resume against a job description.
    
    Returns:
        dict: Complete structured resume representation containing sections, skills, quality analysis, candidate profile, etc.
    """
    filename = os.path.basename(file_path)
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    empty_sections = {sec: "" for sec in CANONICAL_SECTIONS}
    empty_analysis = {
        "score": 0,
        "max_score": 100,
        "categories": {},
        "strengths": [],
        "warnings": [],
        "recommendations": []
    }
    empty_recommendation_plan = {
        "all_recommendations": [],
        "resume_improvements": [],
        "job_improvements": [],
        "has_job_recommendations": False,
        "counts": {"high": 0, "medium": 0, "low": 0, "total": 0}
    }
    empty_evidence_analysis = {
        "evidence_strength_score": 0,
        "skill_evidence_list": [],
        "job_evidence_matrix": [],
        "evidence_gaps": [],
        "keyword_coverage": {"matched": 0, "total": 0, "percentage": 0},
        "has_job_description": False
    }
    empty_candidate_profile = {
        "domains": [],
        "probable_roles": [],
        "skills": [],
        "categorized_skills": {"technical_skills": [], "tools": [], "soft_skills": [], "domain_competencies": []},
        "years_of_experience": 0.0,
        "has_education": False,
        "has_certifications": False,
        "has_summary": False,
        "summary_snippet": ""
    }

    result = {
        "filename": filename,
        "file_type": ext,
        "page_count": None,
        "raw_text": "",
        "cleaned_text": "",
        "has_text": False,
        "sections": empty_sections,
        "skills": [],
        "analysis": empty_analysis,
        "candidate_profile": empty_candidate_profile,
        "job_match": {"has_job_description": False},
        "evidence_analysis": empty_evidence_analysis,
        "recommendation_plan": empty_recommendation_plan,
        "error": None
    }

    if not os.path.exists(file_path):
        result["error"] = f"File not found: {filename}"
        return result

    if ext == "pdf":
        parse_res = parse_pdf(file_path)
    elif ext == "docx":
        parse_res = parse_docx(file_path)
    else:
        result["error"] = f"Unsupported file extension '.{ext}'."
        return result

    result["page_count"] = parse_res.get("page_count")
    result["raw_text"] = parse_res.get("raw_text", "")
    result["error"] = parse_res.get("error")

    if not result["error"]:
        cleaned = clean_text(result["raw_text"])
        result["cleaned_text"] = cleaned
        result["has_text"] = bool(cleaned.strip())
        
        if result["has_text"]:
            result["sections"] = detect_sections(cleaned)
            # Extract skills from the detected 'skills' section or body text
            raw_skills_text = result["sections"].get("skills", "") or cleaned
            result["skills"] = extract_skills(raw_skills_text)
            
            # Build normalized CandidateProfile (Phase 11 AI/NLP Intelligence)
            result["candidate_profile"] = build_candidate_profile(result)

            # Analyze resume quality & compute explainable scores
            result["analysis"] = analyze_resume(result)

            # Match against job description if provided
            if job_description:
                result["job_match"] = match_resume_to_job(result, job_description)

            # Analyze skill locations, evidence strength, matrix & gaps (Phase 8)
            result["evidence_analysis"] = analyze_resume_evidence(result, result.get("job_match"))

            # Generate prioritized actionable improvement plan (Phase 7 & 8)
            result["recommendation_plan"] = generate_recommendations(
                result["analysis"],
                result.get("job_match"),
                result.get("evidence_analysis")
            )

    return result
