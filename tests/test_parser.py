import os
import shutil
import tempfile
import fitz
import docx
import pytest

from services.text_processor import clean_text
from services.resume_parser import parse_pdf, parse_docx, parse_resume


@pytest.fixture
def temp_dir():
    dirpath = tempfile.mkdtemp()
    yield dirpath
    shutil.rmtree(dirpath)


def create_sample_pdf(filepath: str, pages_text: list):
    """Helper utility to generate a real PDF file for testing using PyMuPDF."""
    doc = fitz.open()
    for text in pages_text:
        page = doc.new_page()
        if text:
            page.insert_text((50, 50), text)
    doc.save(filepath)
    doc.close()


def create_sample_docx(filepath: str, paragraphs: list, table_data: list = None):
    """Helper utility to generate a real DOCX file for testing using python-docx."""
    doc = docx.Document()
    for p_text in paragraphs:
        doc.add_paragraph(p_text)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for row_idx, row in enumerate(table_data):
            for col_idx, cell_text in enumerate(row):
                table.cell(row_idx, col_idx).text = cell_text
    doc.save(filepath)


def test_text_processor_cleaning():
    """Verify whitespace normalization, line ending conversion, and gap reduction."""
    raw = "  John Doe  \r\n\r\nSoftware Engineer \t  \n\n\n\nSkills:\nPython,   Flask   \n\n\n"
    cleaned = clean_text(raw)
    assert cleaned == "John Doe\n\nSoftware Engineer\n\nSkills:\nPython, Flask"


def test_pdf_single_page_extraction(temp_dir):
    """Verify single-page PDF text extraction."""
    pdf_path = os.path.join(temp_dir, "single_page.pdf")
    create_sample_pdf(pdf_path, ["John Doe - Software Engineer"])

    res = parse_pdf(pdf_path)
    assert res["error"] is None
    assert res["page_count"] == 1
    assert "John Doe" in res["raw_text"]

    resume_res = parse_resume(pdf_path)
    assert resume_res["file_type"] == "pdf"
    assert resume_res["has_text"] is True
    assert "John Doe" in resume_res["cleaned_text"]


def test_pdf_multipage_extraction(temp_dir):
    """Verify multi-page PDF text extraction and page counting."""
    pdf_path = os.path.join(temp_dir, "multipage.pdf")
    create_sample_pdf(pdf_path, [
        "Page 1: Experience at Company A",
        "Page 2: Education and Skills"
    ])

    res = parse_pdf(pdf_path)
    assert res["error"] is None
    assert res["page_count"] == 2
    assert "Page 1" in res["raw_text"]
    assert "Page 2" in res["raw_text"]


def test_docx_paragraph_extraction(temp_dir):
    """Verify DOCX paragraph extraction."""
    docx_path = os.path.join(temp_dir, "sample.docx")
    create_sample_docx(docx_path, ["Jane Smith", "Senior Developer", "5 years experience"])

    res = parse_docx(docx_path)
    assert res["error"] is None
    assert "Jane Smith" in res["raw_text"]
    assert "Senior Developer" in res["raw_text"]


def test_docx_table_extraction(temp_dir):
    """Verify DOCX table text extraction."""
    docx_path = os.path.join(temp_dir, "table_sample.docx")
    table_data = [
        ["Company", "Role", "Years"],
        ["Acme Corp", "Backend Dev", "2020-2023"]
    ]
    create_sample_docx(docx_path, ["Work History:"], table_data)

    res = parse_docx(docx_path)
    assert res["error"] is None
    assert "Work History:" in res["raw_text"]
    assert "Acme Corp" in res["raw_text"]
    assert "Backend Dev" in res["raw_text"]


def test_empty_or_no_text_document(temp_dir):
    """Verify handling of empty PDF without text content."""
    pdf_path = os.path.join(temp_dir, "empty.pdf")
    create_sample_pdf(pdf_path, [""])

    res = parse_resume(pdf_path)
    assert res["error"] is None
    assert res["has_text"] is False
    assert res["cleaned_text"] == ""


def test_corrupt_pdf_handling(temp_dir):
    """Verify graceful error handling for corrupt PDF file."""
    corrupt_pdf = os.path.join(temp_dir, "corrupt.pdf")
    with open(corrupt_pdf, "wb") as f:
        f.write(b"NOT A REAL PDF FILE CONTENT")

    res = parse_resume(corrupt_pdf)
    assert res["error"] is not None
    assert "Failed to parse PDF" in res["error"]
    assert res["has_text"] is False


def test_corrupt_docx_handling(temp_dir):
    """Verify graceful error handling for corrupt DOCX file."""
    corrupt_docx = os.path.join(temp_dir, "corrupt.docx")
    with open(corrupt_docx, "wb") as f:
        f.write(b"NOT A REAL DOCX ZIP ARCHIVE")

    res = parse_resume(corrupt_docx)
    assert res["error"] is not None
    assert "Failed to parse DOCX" in res["error"]
    assert res["has_text"] is False


def test_unsupported_file_extension(temp_dir):
    """Verify rejection of unsupported file formats."""
    txt_path = os.path.join(temp_dir, "file.txt")
    with open(txt_path, "w") as f:
        f.write("text content")

    res = parse_resume(txt_path)
    assert res["error"] is not None
    assert "Unsupported file extension" in res["error"]
